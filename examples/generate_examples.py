"""Generate synthetic demo files with fictional PII for testing scrubfile."""

import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from pathlib import Path

OUT = Path(__file__).parent


def make_employee_record_pdf():
    """Text-based PDF: HR employee record with varied PII."""
    doc = fitz.open()
    p = doc.new_page()
    y = 72

    def w(text, size=12, bold=False):
        nonlocal y
        p.insert_text(
            (72, y), text, fontsize=size, fontname="helv" if not bold else "hebo"
        )
        y += size + 8

    w("CONFIDENTIAL - EMPLOYEE RECORD", 18, bold=True)
    y += 10
    w("Full Name: James Robert Mitchell")
    w("Date of Birth: March 15, 1988")
    w("Social Security Number: 287-65-4321")
    w("Phone: (555) 867-5309")
    w("Email: james.mitchell@globalcorp.com")
    w("Home Address: 4521 Maple Avenue, Sunnyvale, CA 94086")
    y += 10
    w("Emergency Contact: Laura Mitchell, 555-234-8901")
    w("Relationship: Spouse")
    y += 10
    w("Bank Account (Direct Deposit): Chase, Routing 021000021, Acct 9876543210")
    w("Annual Salary: $142,500")
    y += 10
    w("Notes: James joined the ML Platform team on 2023-01-15.")
    w("Manager: Sarah Chen (sarah.chen@globalcorp.com)")
    w("Previous employer: Meridian Technologies Inc.")

    doc.set_metadata({"author": "HR System v3.2", "title": "Employee Record - Mitchell"})
    doc.save(str(OUT / "sample_employee_record.pdf"))
    doc.close()
    print("Created sample_employee_record.pdf")


def make_medical_form_pdf():
    """Text-based PDF: Medical intake form."""
    doc = fitz.open()
    p = doc.new_page()
    y = 72

    def w(text, size=11):
        nonlocal y
        p.insert_text((72, y), text, fontsize=size, fontname="helv")
        y += size + 7

    w("PATIENT INTAKE FORM", 16)
    y += 8
    w("Patient Name: Jane Elizabeth Morrison")
    w("Date of Birth: 07/22/1985")
    w("SSN: 432-10-8765")
    w("Insurance ID: BCBS-IL-9928374650")
    w("Phone: 312-555-0147")
    w("Email: jane.morrison@protonmail.com")
    w("Address: 789 Oak Boulevard, Apt 4B, Chicago, IL 60614")
    y += 8
    w("Emergency Contact: Robert Morrison (husband), 312.555.0298")
    y += 8
    w("Primary Care Physician: Dr. Michael Patel")
    w("Allergies: Penicillin, Sulfa drugs")
    w("Current Medications: Lisinopril 10mg, Metformin 500mg")
    y += 8
    w("Reason for Visit: Annual physical examination")
    w("Insurance Group #: GRP-445521")
    w("Driver's License: D400-8293-1247 (Illinois)")
    w("Preferred Pharmacy: CVS, 1200 N State St, Chicago IL")

    doc.set_metadata({"author": "MedChart EHR", "title": "Patient Intake - Morrison"})
    doc.save(str(OUT / "sample_medical_form.pdf"))
    doc.close()
    print("Created sample_medical_form.pdf")


def make_contract_docx():
    """DOCX: Employment contract with PII throughout."""
    doc = Document()
    doc.core_properties.author = "Legal Department"
    doc.core_properties.title = "Employment Agreement"

    doc.add_heading("EMPLOYMENT AGREEMENT", level=1)
    doc.add_paragraph(
        'This Employment Agreement ("Agreement") is entered into as of March 1, 2026, '
        'by and between TechVentures Inc., a Delaware corporation ("Company"), and '
        'Maria Sofia Garcia ("Employee").'
    )

    doc.add_heading("1. EMPLOYEE INFORMATION", level=2)
    doc.add_paragraph("Full Legal Name: Maria Sofia Garcia")
    doc.add_paragraph("Social Security Number: 612-73-4589")
    doc.add_paragraph("Date of Birth: September 8, 1990")
    doc.add_paragraph("Address: 2847 Sunset Boulevard, Apt 12C, Los Angeles, CA 90028")
    doc.add_paragraph("Phone: (323) 555-4178")
    doc.add_paragraph("Email: maria.garcia@techventures.com")
    doc.add_paragraph("Personal Email: mgarcia1990@gmail.com")

    doc.add_heading("2. COMPENSATION", level=2)
    doc.add_paragraph("Base Salary: $185,000.00 per year")
    doc.add_paragraph(
        "Direct Deposit: Wells Fargo, Routing #121000248, Account #7834291056"
    )
    doc.add_paragraph("Sign-on Bonus: $25,000 (payable within 30 days of start date)")

    doc.add_heading("3. EMERGENCY CONTACT", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    cells = [
        ("Name", "Carlos Garcia"),
        ("Relationship", "Father"),
        ("Phone", "323-555-8834"),
        ("Email", "c.garcia47@yahoo.com"),
    ]
    for i, (key, val) in enumerate(cells):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "Employee acknowledges receipt of company laptop (Serial: SN-MBP-2026-38471) "
        "and building access card (Badge ID: ACC-LA-0394)."
    )

    doc.save(str(OUT / "sample_contract.docx"))
    print("Created sample_contract.docx")


def make_id_card_image():
    """JPEG image simulating a scanned ID card (needs OCR)."""
    img = Image.new("RGB", (600, 380), color=(230, 235, 240))
    draw = ImageDraw.Draw(img)

    # Card border
    draw.rectangle([10, 10, 590, 370], outline=(0, 50, 100), width=3)

    # Header
    draw.rectangle([10, 10, 590, 60], fill=(0, 50, 100))
    draw.text((150, 20), "CORPORATE ID CARD", fill="white")

    # Photo placeholder
    draw.rectangle([30, 80, 170, 250], fill=(180, 180, 180), outline=(100, 100, 100))
    draw.text((65, 155), "PHOTO", fill=(120, 120, 120))

    # Info
    info = [
        ("Name: Raj Kumar Patel", 200, 90),
        ("Employee ID: EMP-78432", 200, 120),
        ("Department: Engineering", 200, 150),
        ("DOB: 11/03/1988", 200, 180),
        ("SSN: 521-84-3967", 200, 210),
        ("Badge #: B-2026-0193", 200, 240),
        ("Phone: 650-555-7712", 200, 270),
        ("raj.patel@techcorp.io", 200, 300),
    ]

    for text, x, y_pos in info:
        draw.text((x, y_pos), text, fill=(20, 20, 20))

    draw.text((200, 340), "Valid: 01/2026 - 12/2026", fill=(100, 100, 100))

    img.save(str(OUT / "sample_id_card.jpg"), quality=85)
    print("Created sample_id_card.jpg")


def make_scanned_letter_image():
    """PNG image simulating a scanned business letter (needs OCR)."""
    img = Image.new("RGB", (850, 1100), color=(252, 250, 245))  # slightly off-white
    draw = ImageDraw.Draw(img)

    lines = [
        ("GlobalTech Solutions Inc.", 60, 80, 20),
        ("1500 Innovation Drive, Suite 300", 60, 110, 12),
        ("San Jose, CA 95134", 60, 130, 12),
        ("", 0, 0, 0),
        ("March 29, 2026", 60, 170, 12),
        ("", 0, 0, 0),
        ("Mr. David Chen", 60, 210, 12),
        ("4782 Willow Creek Lane", 60, 230, 12),
        ("Portland, OR 97201", 60, 250, 12),
        ("", 0, 0, 0),
        ("Dear Mr. Chen,", 60, 290, 12),
        ("", 0, 0, 0),
        ("We are pleased to confirm your employment offer.", 60, 330, 12),
        ("Your employee ID will be EMP-2026-0847.", 60, 355, 12),
        ("Please provide your SSN (currently on file as", 60, 380, 12),
        ("378-92-6154) for tax processing.", 60, 405, 12),
        ("", 0, 0, 0),
        ("Your starting salary will be $165,000 annually.", 60, 445, 12),
        ("Direct deposit to account ending in 4829.", 60, 470, 12),
        ("Contact HR at hr@globaltech.com or 408-555-9321.", 60, 495, 12),
        ("", 0, 0, 0),
        ("Sincerely,", 60, 535, 12),
        ("Amanda Rodriguez", 60, 575, 14),
        ("VP of Human Resources", 60, 600, 12),
        ("amanda.rodriguez@globaltech.com", 60, 625, 12),
    ]

    for text, x, y_pos, size in lines:
        if text:
            draw.text((x, y_pos), text, fill=(30, 30, 30))

    img.save(str(OUT / "sample_scanned_letter.png"))
    print("Created sample_scanned_letter.png")


if __name__ == "__main__":
    make_employee_record_pdf()
    make_medical_form_pdf()
    make_contract_docx()
    make_id_card_image()
    make_scanned_letter_image()
    print("\nAll example files generated in:", OUT)
