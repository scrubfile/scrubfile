"""Redactor — Local PII redaction tool for documents."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from redactor.utils import (
    expand_term_variants,
    expand_thorough_variants,
    resolve_output_path,
    validate_input_file,
)


@dataclass
class RedactionResult:
    """Unified result for any redaction operation."""

    input_path: str
    output_path: str
    total_redactions: int
    pages_affected: int = 0
    terms_found: dict[str, int] = field(default_factory=dict)
    metadata_cleared: bool = False


_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"}


def redact(
    file_path: str | Path,
    terms: list[str] | None = None,
    output: str | Path | None = None,
    ocr_engine: str = "easyocr",
    auto: bool = False,
    threshold: float = 0.7,
    entity_types: list[str] | None = None,
    preview: bool = False,
    thorough: bool = False,
) -> RedactionResult:
    """Redact PII terms from a document.

    This is the primary public API. It auto-detects the file type and
    routes to the appropriate handler.

    Args:
        file_path: Path to the document to redact.
        terms: List of PII strings to search for and redact.
        output: Optional output path. Defaults to <input>_redacted_<timestamp>.<ext>.
        ocr_engine: OCR engine for image files ("easyocr" or "tesseract").
        auto: If True, auto-detect PII using NLP (Presidio + spaCy).
        threshold: Confidence threshold for auto-detection (0.0-1.0). Default 0.7.
        entity_types: Entity types to detect in auto mode (e.g. ["PERSON", "US_SSN"]).
        preview: If True with auto=True, return detections without redacting.
        thorough: If True, also redact individual name components and fragments
            to prevent residual inference attacks. E.g., "John Doe" also redacts
            "John", "Doe", "J. Doe". Increases false positives.

    Returns:
        RedactionResult with details about what was redacted.

    Raises:
        FileNotFoundError: If the input file doesn't exist.
        ValueError: If the file type is unsupported or file is too large.
    """
    input_path = validate_input_file(file_path)
    output_path = resolve_output_path(input_path, output)

    # Choose expansion strategy
    _expand = expand_thorough_variants if thorough else expand_term_variants

    # Auto-detect PII if requested
    if auto:
        auto_terms = _auto_detect_terms(input_path, threshold, entity_types, ocr_engine)
        if terms:
            all_terms = list(set(_expand(terms)) | set(auto_terms))
        else:
            all_terms = auto_terms
        # In thorough mode, also expand auto-detected terms
        if thorough:
            all_terms = expand_thorough_variants(all_terms)

        if preview:
            return RedactionResult(
                input_path=str(input_path),
                output_path=str(output_path),
                total_redactions=len(all_terms),
                terms_found={t: 1 for t in all_terms},
                metadata_cleared=False,
            )
        terms = all_terms
    elif terms:
        terms = _expand(terms)
    else:
        raise ValueError("Provide terms to redact, or use auto=True for auto-detection.")

    suffix = input_path.suffix.lower()

    if suffix == ".pdf":
        from redactor.pdf import redact_pdf

        r = redact_pdf(input_path, output_path, terms)
        return RedactionResult(
            input_path=r.input_path,
            output_path=r.output_path,
            total_redactions=r.total_redactions,
            pages_affected=r.pages_affected,
            terms_found=r.terms_found,
            metadata_cleared=r.metadata_cleared,
        )

    elif suffix in _IMAGE_EXTENSIONS:
        from redactor.image import redact_image

        r = redact_image(input_path, output_path, terms, ocr_engine=ocr_engine)
        return RedactionResult(
            input_path=r.input_path,
            output_path=r.output_path,
            total_redactions=r.total_redactions,
            terms_found=r.terms_found,
            metadata_cleared=r.metadata_cleared,
        )

    elif suffix == ".docx":
        from redactor.docx_redactor import redact_docx

        r = redact_docx(input_path, output_path, terms)
        return RedactionResult(
            input_path=r.input_path,
            output_path=r.output_path,
            total_redactions=r.total_redactions,
            terms_found=r.terms_found,
            metadata_cleared=r.metadata_cleared,
        )

    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _auto_detect_terms(
    input_path: Path,
    threshold: float,
    entity_types: list[str] | None,
    ocr_engine: str,
) -> list[str]:
    """Extract text from document and detect PII terms automatically."""
    from redactor.detector import detect_pii

    suffix = input_path.suffix.lower()
    text = _extract_text(input_path, suffix, ocr_engine)

    detections = detect_pii(
        text=text,
        threshold=threshold,
        entity_types=entity_types,
    )

    # Deduplicate detected terms
    seen: set[str] = set()
    terms: list[str] = []
    for d in detections:
        t = d.text.strip()
        if t and t not in seen:
            seen.add(t)
            terms.append(t)

    return terms


def _extract_text(input_path: Path, suffix: str, ocr_engine: str) -> str:
    """Extract text from a document for PII detection."""
    if suffix == ".pdf":
        from redactor.pdf import extract_text
        return extract_text(input_path)

    elif suffix in _IMAGE_EXTENSIONS:
        from redactor.ocr import get_ocr_provider
        from PIL import Image

        provider = get_ocr_provider(ocr_engine)
        image = Image.open(str(input_path))
        results = provider.extract(image)
        return " ".join(r.text for r in results)

    elif suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(input_path))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)

    else:
        raise ValueError(f"Cannot extract text from: {suffix}")
