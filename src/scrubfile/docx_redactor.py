"""DOCX redaction engine using python-docx."""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


@dataclass
class DocxRedactionResult:
    """Result of a DOCX redaction operation."""

    input_path: str
    output_path: str
    total_redactions: int
    terms_found: dict[str, int] = field(default_factory=dict)
    metadata_cleared: bool = False


# Unicode full block character for redaction
_REDACT_CHAR = "\u2588"


def redact_docx(
    input_path: Path,
    output_path: Path,
    terms: list[str],
) -> DocxRedactionResult:
    """Redact PII terms from a DOCX file.

    Searches paragraphs, tables, headers, and footers.
    Replaces matched text with block characters (████).
    Clears document properties (author, company, etc.).
    """
    doc = Document(str(input_path))

    total_redactions = 0
    terms_found: dict[str, int] = {}

    # Process all text containers
    for paragraph in _all_paragraphs(doc):
        count = _redact_paragraph(paragraph, terms)
        for term, n in count.items():
            terms_found[term] = terms_found.get(term, 0) + n
            total_redactions += n

    # Clear document properties
    _clear_properties(doc)

    doc.save(str(output_path))
    os.chmod(str(output_path), 0o600)

    return DocxRedactionResult(
        input_path=str(input_path),
        output_path=str(output_path),
        total_redactions=total_redactions,
        terms_found=terms_found,
        metadata_cleared=True,
    )


def _all_paragraphs(doc: Document):
    """Yield all paragraphs from body, tables, headers, and footers."""
    # Body paragraphs
    yield from doc.paragraphs

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs

    # Header and footer paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _redact_paragraph(paragraph, terms: list[str]) -> dict[str, int]:
    """Redact terms within a paragraph, handling cross-run spans.

    Word splits text into "runs" at formatting boundaries. A term like
    "John Doe" might be split across runs as ["Jo", "hn Doe"]. This
    function concatenates all runs, finds matches in the full text,
    then maps redactions back to individual runs.
    """
    counts: dict[str, int] = {}

    # Get full paragraph text by joining runs
    runs = paragraph.runs
    if not runs:
        return counts

    full_text = "".join(r.text for r in runs)
    if not full_text.strip():
        return counts

    # Build a map: character index in full_text → (run_index, char_index_within_run)
    char_map: list[tuple[int, int]] = []
    for run_idx, run in enumerate(runs):
        for char_idx in range(len(run.text)):
            char_map.append((run_idx, char_idx))

    # Find all term matches (case-insensitive)
    for term in terms:
        pattern = re.compile(re.escape(term), re.IGNORECASE)
        for match in pattern.finditer(full_text):
            start, end = match.start(), match.end()
            counts[term] = counts.get(term, 0) + 1

            # Replace matched characters with redaction blocks
            replacement = _REDACT_CHAR * (end - start)
            _replace_in_runs(runs, char_map, start, end, replacement)

    return counts


def _replace_in_runs(runs, char_map, start: int, end: int, replacement: str):
    """Replace characters from start to end across runs."""
    # Group characters by run
    affected_runs: dict[int, list[tuple[int, str]]] = {}
    for i, repl_char in zip(range(start, end), replacement):
        run_idx, char_idx = char_map[i]
        affected_runs.setdefault(run_idx, []).append((char_idx, repl_char))

    # Apply replacements to each affected run
    for run_idx, replacements in affected_runs.items():
        text = list(runs[run_idx].text)
        for char_idx, repl_char in replacements:
            text[char_idx] = repl_char
        runs[run_idx].text = "".join(text)


def _clear_properties(doc: Document):
    """Clear document core properties (author, company, etc.)."""
    props = doc.core_properties
    props.author = ""
    props.category = ""
    props.comments = ""
    props.content_status = ""
    props.keywords = ""
    props.last_modified_by = ""
    props.subject = ""
    props.title = ""
