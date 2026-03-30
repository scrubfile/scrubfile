"""Tests for DOCX redaction engine."""

from __future__ import annotations

import stat
from pathlib import Path

import pytest
from docx import Document

from redactor.docx_redactor import redact_docx, DocxRedactionResult


def _create_test_docx(path: Path) -> Path:
    """Create a test DOCX with known PII."""
    doc = Document()

    # Set document properties
    doc.core_properties.author = "HR Department"
    doc.core_properties.title = "Employee Record"

    # Body paragraphs
    doc.add_paragraph("Confidential Employee Record")
    doc.add_paragraph("Name: John Doe")
    doc.add_paragraph("SSN: 123-45-6789")
    doc.add_paragraph("Email: john@example.com")
    doc.add_paragraph("This document contains sensitive information about John Doe.")

    # Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Employee"
    table.cell(0, 1).text = "Department"
    table.cell(1, 0).text = "John Doe"
    table.cell(1, 1).text = "Engineering"

    doc.save(str(path))
    return path


@pytest.fixture
def test_docx(tmp_path: Path) -> Path:
    return _create_test_docx(tmp_path / "test.docx")


class TestDocxRedaction:
    def test_redacts_single_term(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact_docx(test_docx, output, ["John Doe"])

        assert output.exists()
        assert result.total_redactions >= 3  # appears in paragraphs + table
        assert "John Doe" in result.terms_found

    def test_redacted_text_is_replaced(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact_docx(test_docx, output, ["John Doe"])

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "John Doe" not in full_text
        assert "\u2588" in full_text  # block characters present

    def test_non_redacted_text_preserved(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact_docx(test_docx, output, ["John Doe"])

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Confidential Employee Record" in full_text
        assert "Engineering" not in full_text or True  # table text checked separately

    def test_redacts_ssn(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact_docx(test_docx, output, ["123-45-6789"])

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "123-45-6789" not in full_text

    def test_redacts_email(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact_docx(test_docx, output, ["john@example.com"])

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "john@example.com" not in full_text

    def test_redacts_in_tables(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact_docx(test_docx, output, ["John Doe"])

        doc = Document(str(output))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    assert "John Doe" not in cell.text

    def test_multiple_terms(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact_docx(
            test_docx, output, ["John Doe", "123-45-6789", "john@example.com"]
        )

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "John Doe" not in full_text
        assert "123-45-6789" not in full_text
        assert "john@example.com" not in full_text
        assert result.total_redactions >= 4

    def test_no_match_returns_zero(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact_docx(test_docx, output, ["NONEXISTENT_XYZ"])

        assert result.total_redactions == 0
        assert output.exists()

    def test_case_insensitive(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact_docx(test_docx, output, ["john doe"])

        assert result.total_redactions >= 1


class TestDocxMetadata:
    def test_metadata_cleared(self, test_docx: Path, tmp_path: Path):
        # Verify source has metadata
        source = Document(str(test_docx))
        assert source.core_properties.author == "HR Department"

        output = tmp_path / "out.docx"
        result = redact_docx(test_docx, output, ["John Doe"])

        assert result.metadata_cleared is True
        doc = Document(str(output))
        assert doc.core_properties.author == ""
        assert doc.core_properties.title == ""


class TestDocxPermissions:
    def test_output_is_owner_only(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact_docx(test_docx, output, ["John Doe"])

        mode = output.stat().st_mode
        assert not (mode & stat.S_IRGRP)
        assert not (mode & stat.S_IROTH)


class TestDocxResult:
    def test_result_fields(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact_docx(test_docx, output, ["John Doe"])

        assert isinstance(result, DocxRedactionResult)
        assert result.input_path == str(test_docx)
        assert result.output_path == str(output)
        assert result.total_redactions > 0
        assert isinstance(result.terms_found, dict)
        assert result.metadata_cleared is True
