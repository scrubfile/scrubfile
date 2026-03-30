"""Tests for auto-detection integrated with redaction."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from scrubfile import RedactionResult, redact
from scrubfile.pdf import extract_text


@pytest.fixture
def pii_pdf(tmp_path: Path) -> Path:
    """Create a PDF with varied PII for auto-detection."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Employee Record", fontsize=16, fontname="helv")
    page.insert_text((72, 140), "Name: John Smith", fontsize=12, fontname="helv")
    page.insert_text((72, 170), "Email: john.smith@example.com", fontsize=12, fontname="helv")
    page.insert_text((72, 200), "Phone: 555-123-4567", fontsize=12, fontname="helv")
    page.insert_text((72, 230), "SSN: 123-45-6789", fontsize=12, fontname="helv")
    doc.set_metadata({"author": "HR System"})
    doc.save(str(tmp_path / "pii.pdf"))
    doc.close()
    return tmp_path / "pii.pdf"


@pytest.fixture
def pii_docx(tmp_path: Path) -> Path:
    """Create a DOCX with varied PII for auto-detection."""
    doc = Document()
    doc.add_paragraph("Employee: John Smith")
    doc.add_paragraph("Email: john.smith@example.com")
    doc.add_paragraph("SSN: 123-45-6789")
    doc.save(str(tmp_path / "pii.docx"))
    return tmp_path / "pii.docx"


class TestAutoRedactPdf:
    def test_auto_detects_and_redacts(self, pii_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        result = redact(pii_pdf, auto=True, threshold=0.3, output=output)

        assert isinstance(result, RedactionResult)
        assert result.total_redactions >= 1
        assert output.exists()

    def test_auto_removes_detected_pii(self, pii_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        redact(pii_pdf, auto=True, threshold=0.3, output=output)

        text = extract_text(output)
        # At least email should be detected and removed
        assert "john.smith@example.com" not in text

    def test_preview_does_not_modify(self, pii_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        result = redact(pii_pdf, auto=True, preview=True, threshold=0.3, output=output)

        assert result.metadata_cleared is False
        # The output file should NOT be created in preview mode
        # (the original is not modified)
        assert len(result.terms_found) >= 1

    def test_auto_plus_explicit_terms(self, pii_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        result = redact(
            pii_pdf,
            terms=["Employee Record"],
            auto=True,
            threshold=0.3,
            output=output,
        )

        text = extract_text(output)
        assert "Employee Record" not in text
        assert result.total_redactions >= 2

    def test_high_threshold_fewer_detections(self, pii_pdf: Path, tmp_path: Path):
        out_low = tmp_path / "low.pdf"
        out_high = tmp_path / "high.pdf"

        r_low = redact(pii_pdf, auto=True, threshold=0.1, output=out_low)
        r_high = redact(pii_pdf, auto=True, threshold=0.95, output=out_high)

        assert r_low.total_redactions >= r_high.total_redactions

    def test_entity_type_filter(self, pii_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        result = redact(
            pii_pdf,
            auto=True,
            threshold=0.3,
            entity_types=["EMAIL_ADDRESS"],
            output=output,
        )

        text = extract_text(output)
        assert "john.smith@example.com" not in text


class TestAutoRedactDocx:
    def test_auto_detects_in_docx(self, pii_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact(pii_docx, auto=True, threshold=0.3, output=output)

        assert isinstance(result, RedactionResult)
        assert result.total_redactions >= 1

    def test_auto_removes_email_from_docx(self, pii_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        redact(pii_docx, auto=True, threshold=0.3, output=output)

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "john.smith@example.com" not in full_text


class TestNoTermsNoAuto:
    def test_raises_without_terms_or_auto(self, pii_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        with pytest.raises(ValueError, match="Provide terms"):
            redact(pii_pdf, output=output)
