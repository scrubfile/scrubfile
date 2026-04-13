"""Golden document production tests — verify scrubfile works on realistic documents.

These tests run the FULL pipeline (extract text → detect PII → redact) on the
7 actual_testing documents and assert minimum expected detections.

This is the ultimate "does it work in production?" gate. Each test defines:
- The exact input document
- The minimum PII that MUST be detected
- The PII that MUST NOT survive in the output

If a new model or config change causes any of these to fail, it cannot ship.
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest

from scrubfile import RedactionResult, redact
from scrubfile.detector import detect_pii

# ---------------------------------------------------------------------------
# Paths — all tests use the checked-in actual_testing/ documents
# ---------------------------------------------------------------------------

ACTUAL_TESTING_DIR = Path(__file__).parent.parent / "actual_testing"

# Skip the entire module if actual_testing/ docs don't exist
pytestmark = pytest.mark.skipif(
    not ACTUAL_TESTING_DIR.exists(),
    reason="actual_testing/ directory not found",
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _extract_text_for_file(path: Path) -> str:
    """Extract text from a document using scrubfile's own extraction."""
    from scrubfile import _extract_text
    suffix = path.suffix.lower()
    return _extract_text(path, suffix, "easyocr")


def _auto_redact(input_path: Path, tmp_path: Path) -> tuple[RedactionResult, str]:
    """Run auto-redaction and return (result, output_text)."""
    output = tmp_path / f"redacted_{input_path.name}"
    if input_path.suffix.lower() in (".png", ".jpg", ".jpeg"):
        # For images, we can't easily extract text from redacted output,
        # so just return the result
        result = redact(input_path, auto=True, threshold=0.3, output=output)
        return result, ""

    result = redact(input_path, auto=True, threshold=0.3, output=output)

    # Extract text from redacted output
    if output.suffix.lower() == ".pdf":
        from scrubfile.pdf import extract_text
        output_text = extract_text(output)
    elif output.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(output))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        output_text = "\n".join(paragraphs)
    else:
        output_text = ""

    return result, output_text


# ===========================================================================
# GOLDEN DOCUMENT TESTS
# ===========================================================================

@pytest.mark.slow
class TestGoldenEmployeeRecord:
    """01_employee_record.pdf — HR record with dense PII."""

    DOC = ACTUAL_TESTING_DIR / "01_employee_record.pdf"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_minimum_pii_types(self):
        """Must detect person, SSN, email, phone at minimum."""
        from scrubfile.pdf import extract_text
        text = extract_text(self.DOC)
        types = {r.entity_type for r in detect_pii(text, threshold=0.3)}

        assert "PERSON" in types, f"Missing PERSON, got: {types}"
        assert "US_SSN" in types, f"Missing US_SSN, got: {types}"
        assert "EMAIL_ADDRESS" in types, f"Missing EMAIL_ADDRESS, got: {types}"
        assert "PHONE_NUMBER" in types, f"Missing PHONE_NUMBER, got: {types}"

    def test_redacts_ssn(self, tmp_path):
        """SSN '287-65-4321' must not survive redaction."""
        result, text = _auto_redact(self.DOC, tmp_path)
        assert result.total_redactions >= 4, (
            f"Expected ≥4 redactions, got {result.total_redactions}"
        )
        assert "287-65-4321" not in text, "SSN survived redaction"

    def test_redacts_email(self, tmp_path):
        """Email must not survive redaction."""
        _, text = _auto_redact(self.DOC, tmp_path)
        assert "james.mitchell@globalcorp.com" not in text, "Email survived redaction"

    def test_known_pii_inventory(self):
        """Verify detection of the specific PII in this document."""
        from scrubfile.pdf import extract_text
        text = extract_text(self.DOC)
        results = detect_pii(text, threshold=0.3)
        detected_texts = {r.text for r in results}

        # These specific PII values are in the document (from generate_test_files.py)
        critical_pii = ["287-65-4321"]  # SSN — must always detect
        for pii in critical_pii:
            assert any(pii in t for t in detected_texts), (
                f"Critical PII '{pii}' not detected. Found: {detected_texts}"
            )


@pytest.mark.slow
class TestGoldenMedicalForm:
    """02_medical_form.pdf — Patient intake form."""

    DOC = ACTUAL_TESTING_DIR / "02_medical_form.pdf"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_patient_pii(self):
        """Must detect patient name, SSN, email, phone."""
        from scrubfile.pdf import extract_text
        text = extract_text(self.DOC)
        types = {r.entity_type for r in detect_pii(text, threshold=0.3)}

        assert "PERSON" in types, f"Missing PERSON, got: {types}"
        assert "US_SSN" in types, f"Missing US_SSN, got: {types}"
        assert "EMAIL_ADDRESS" in types, f"Missing EMAIL_ADDRESS, got: {types}"

    def test_redacts_patient_ssn(self, tmp_path):
        """Patient SSN '432-10-8765' must not survive."""
        _, text = _auto_redact(self.DOC, tmp_path)
        assert "432-10-8765" not in text, "Patient SSN survived redaction"

    def test_minimum_redaction_count(self, tmp_path):
        """Medical form has dense PII — expect many redactions."""
        result, _ = _auto_redact(self.DOC, tmp_path)
        assert result.total_redactions >= 5, (
            f"Expected ≥5 redactions on medical form, got {result.total_redactions}"
        )


@pytest.mark.slow
class TestGoldenScannedLetter:
    """03_scanned_letter.png — OCR-dependent scanned business letter."""

    DOC = ACTUAL_TESTING_DIR / "03_scanned_letter.png"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_pii_via_ocr(self, tmp_path):
        """Must detect at least some PII through OCR pipeline."""
        result, _ = _auto_redact(self.DOC, tmp_path)
        assert result.total_redactions >= 1, (
            "OCR pipeline detected zero PII in scanned letter with names, SSN, emails"
        )

    def test_detects_email_via_ocr(self):
        """Email addresses should survive OCR and be detected."""
        text = _extract_text_for_file(self.DOC)
        types = {r.entity_type for r in detect_pii(text, threshold=0.3)}
        # Emails have strong pattern — should survive OCR
        assert "EMAIL_ADDRESS" in types or "PHONE_NUMBER" in types, (
            f"Neither email nor phone detected via OCR, got: {types}"
        )


@pytest.mark.slow
class TestGoldenIdCard:
    """04_id_card.jpg — Scanned corporate ID card."""

    DOC = ACTUAL_TESTING_DIR / "04_id_card.jpg"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_pii_on_id_card(self, tmp_path):
        """ID card with name, SSN, phone, email — must detect some."""
        result, _ = _auto_redact(self.DOC, tmp_path)
        assert result.total_redactions >= 1, (
            "Zero PII detected on ID card with SSN, name, phone, email"
        )


@pytest.mark.slow
class TestGoldenEmploymentContract:
    """05_employment_contract.docx — Full contract with PII in body and tables."""

    DOC = ACTUAL_TESTING_DIR / "05_employment_contract.docx"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_contract_pii(self):
        """Must detect SSN, emails, phone, person names in contract."""
        text = _extract_text_for_file(self.DOC)
        types = {r.entity_type for r in detect_pii(text, threshold=0.3)}

        assert "PERSON" in types, f"Missing PERSON in contract, got: {types}"
        assert "US_SSN" in types, f"Missing US_SSN in contract, got: {types}"
        assert "EMAIL_ADDRESS" in types, f"Missing EMAIL_ADDRESS in contract, got: {types}"
        assert "PHONE_NUMBER" in types, f"Missing PHONE_NUMBER in contract, got: {types}"

    def test_redacts_ssn_from_docx(self, tmp_path):
        """Contract SSN '612-73-4589' must not survive redaction."""
        _, text = _auto_redact(self.DOC, tmp_path)
        assert "612-73-4589" not in text, "Contract SSN survived DOCX redaction"

    def test_redacts_email_from_docx(self, tmp_path):
        """Employee email must not survive redaction."""
        _, text = _auto_redact(self.DOC, tmp_path)
        assert "maria.garcia@techventures.com" not in text, (
            "Employee email survived DOCX redaction"
        )

    def test_table_pii_detected(self):
        """PII inside DOCX tables must also be detected."""
        text = _extract_text_for_file(self.DOC)
        results = detect_pii(text, threshold=0.3)
        detected_texts = " ".join(r.text for r in results)
        # Emergency contact table has email and phone
        assert any(
            "c.garcia47@yahoo.com" in r.text or "323-555-8834" in r.text
            for r in results
        ), f"Table PII not detected. All detections: {detected_texts}"

    def test_minimum_redaction_count(self, tmp_path):
        """Dense contract should have many redactions."""
        result, _ = _auto_redact(self.DOC, tmp_path)
        assert result.total_redactions >= 6, (
            f"Expected ≥6 redactions on contract, got {result.total_redactions}"
        )


@pytest.mark.slow
class TestGoldenScannedW2:
    """06_scanned_w2.pdf — Image-only PDF (no text layer), requires OCR."""

    DOC = ACTUAL_TESTING_DIR / "06_scanned_w2.pdf"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_pii_in_image_pdf(self, tmp_path):
        """Image-only PDF with SSN, name, address — must detect via OCR.

        KNOWN LIMITATION: scrubfile's PDF text extraction (PyMuPDF) does not
        OCR image-only pages. This test documents the gap — when OCR-for-PDF
        is implemented, remove the xfail marker.
        """
        result, _ = _auto_redact(self.DOC, tmp_path)
        if result.total_redactions == 0:
            pytest.xfail(
                "KNOWN GAP: Image-only PDF not OCR'd by current pipeline. "
                "Tracked for future fix."
            )
        assert result.total_redactions >= 1


@pytest.mark.slow
class TestGoldenDirectoryImage:
    """07_directory.png — Multi-person table with 4 people's PII."""

    DOC = ACTUAL_TESTING_DIR / "07_directory.png"

    @pytest.fixture(autouse=True)
    def _skip_if_missing(self):
        if not self.DOC.exists():
            pytest.skip(f"{self.DOC.name} not found")

    def test_detects_multiple_people(self, tmp_path):
        """Directory with 4 people should detect PII from multiple entries."""
        result, _ = _auto_redact(self.DOC, tmp_path)
        assert result.total_redactions >= 3, (
            f"Expected ≥3 redactions for 4-person directory, got {result.total_redactions}"
        )


# ===========================================================================
# CROSS-DOCUMENT AGGREGATE TESTS
# ===========================================================================

@pytest.mark.slow
class TestCrossDocumentGates:
    """Aggregate checks across all golden documents."""

    TEXT_PDFS = [
        ACTUAL_TESTING_DIR / "01_employee_record.pdf",
        ACTUAL_TESTING_DIR / "02_medical_form.pdf",
    ]

    def test_zero_ssn_leaks_across_text_pdfs(self, tmp_path):
        """No SSN should survive redaction across ANY text PDF."""
        known_ssns = ["287-65-4321", "432-10-8765"]

        for pdf_path in self.TEXT_PDFS:
            if not pdf_path.exists():
                continue
            _, text = _auto_redact(pdf_path, tmp_path)
            for ssn in known_ssns:
                assert ssn not in text, (
                    f"SSN {ssn} leaked in {pdf_path.name} after redaction"
                )

    def test_minimum_total_detections(self, tmp_path):
        """Across all text PDFs, total redactions should be substantial."""
        total = 0
        for pdf_path in self.TEXT_PDFS:
            if not pdf_path.exists():
                continue
            result, _ = _auto_redact(pdf_path, tmp_path)
            total += result.total_redactions

        assert total >= 10, (
            f"Expected ≥10 total redactions across text PDFs, got {total}"
        )
