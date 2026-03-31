"""Tests for the public API with multiple file formats."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from PIL import Image, ImageDraw

from scrubfile import RedactionResult, redact


@pytest.fixture
def test_png(tmp_path: Path) -> Path:
    img = Image.new("RGB", (400, 200), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 50), "Hello World", fill="black")
    path = tmp_path / "test.png"
    img.save(str(path))
    return path


@pytest.fixture
def test_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Name: John Doe")
    doc.add_paragraph("SSN: 123-45-6789")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


class TestMultiFormatApi:
    def test_redact_pdf(self, golden_pdf: Path, tmp_path: Path):
        output = tmp_path / "out.pdf"
        result = redact(golden_pdf, terms=["John Doe"], output=output)

        assert isinstance(result, RedactionResult)
        assert result.total_redactions >= 3
        assert output.exists()

    @pytest.mark.slow
    def test_redact_image(self, test_png: Path, tmp_path: Path):
        output = tmp_path / "out.png"
        result = redact(test_png, terms=["Hello"], output=output)

        assert isinstance(result, RedactionResult)
        assert output.exists()

    def test_redact_docx(self, test_docx: Path, tmp_path: Path):
        output = tmp_path / "out.docx"
        result = redact(test_docx, terms=["John Doe"], output=output)

        assert isinstance(result, RedactionResult)
        assert result.total_redactions >= 1
        assert output.exists()

    def test_unsupported_extension(self, tmp_path: Path):
        txt = tmp_path / "test.txt"
        txt.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported file type"):
            redact(txt, terms=["hello"])

    @pytest.mark.slow
    def test_ocr_engine_parameter(self, test_png: Path, tmp_path: Path):
        output = tmp_path / "out.png"
        # Should work with default easyocr
        result = redact(test_png, terms=["Hello"], output=output, ocr_engine="easyocr")
        assert isinstance(result, RedactionResult)

    def test_unified_result_type(self, golden_pdf: Path, test_docx: Path, tmp_path: Path):
        """All file types return the same RedactionResult type."""
        pdf_out = tmp_path / "out.pdf"
        docx_out = tmp_path / "out.docx"

        r1 = redact(golden_pdf, terms=["John Doe"], output=pdf_out)
        r2 = redact(test_docx, terms=["John Doe"], output=docx_out)

        assert type(r1) == type(r2) == RedactionResult
